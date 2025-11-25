"""
文档生成器模块
用于生成分析报告和优化后的简历文档（PDF 和 Word 格式）
"""

from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
from typing import Dict, Any, List
import json


class ReportGenerator:
    """分析报告生成器"""

    def __init__(self):
        pass

    def generate_analysis_report_word(
        self,
        analysis_data: Dict[str, Any],
        recommendations: str = "",
        chat_history: List[Dict] = None
    ) -> BytesIO:
        """
        生成 Word 格式的分析报告

        Args:
            analysis_data: 简历分析数据
            recommendations: AI 推荐内容
            chat_history: 对话历史

        Returns:
            BytesIO: Word 文档的字节流
        """
        doc = Document()

        # 设置中文字体
        doc.styles['Normal'].font.name = 'SimSun'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        # 标题
        title = doc.add_heading('简历分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 生成时间
        time_para = doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}')
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # 空行

        # 1. 简历分析摘要
        doc.add_heading('一、简历分析摘要', 1)

        # 基本信息表格
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'

        table.cell(0, 0).text = '工作年限'
        table.cell(0, 1).text = f"{analysis_data.get('experience_years', 0)} 年"

        table.cell(1, 0).text = '学历水平'
        table.cell(1, 1).text = analysis_data.get('education', '未知')

        table.cell(2, 0).text = '期望职位'
        table.cell(2, 1).text = analysis_data.get('desired_position', '未提及')

        table.cell(3, 0).text = '期望薪资'
        table.cell(3, 1).text = analysis_data.get('desired_salary', '未提及')

        table.cell(4, 0).text = '技能数量'
        skills = analysis_data.get('skills', [])
        table.cell(4, 1).text = f"{len(skills)} 项"

        table.cell(5, 0).text = '核心优势'
        key_strengths = analysis_data.get('key_strengths', [])
        table.cell(5, 1).text = f"{len(key_strengths)} 项"

        doc.add_paragraph()

        # 2. 技能清单
        doc.add_heading('二、技能清单', 1)
        if skills:
            skills_para = doc.add_paragraph()
            for skill in skills:
                skills_para.add_run(f'• {skill}\n')
        else:
            doc.add_paragraph('未识别到具体技能')

        doc.add_paragraph()

        # 3. 核心优势
        doc.add_heading('三、核心优势', 1)
        if key_strengths:
            for i, strength in enumerate(key_strengths, 1):
                doc.add_paragraph(f'{i}. {strength}', style='List Number')
        else:
            doc.add_paragraph('未识别到核心优势')

        doc.add_paragraph()

        # 4. 工作经历
        doc.add_heading('四、工作经历摘要', 1)
        work_exp = analysis_data.get('work_experience', '未提供工作经历')
        doc.add_paragraph(work_exp)

        doc.add_paragraph()

        # 5. AI 职位推荐
        if recommendations:
            doc.add_heading('五、AI 职位推荐', 1)
            doc.add_paragraph(recommendations)
            doc.add_paragraph()

        # 6. 对话记录（可选）
        if chat_history and len(chat_history) > 0:
            doc.add_heading('六、AI 对话记录', 1)
            for msg in chat_history[:10]:  # 只显示前10条
                role = '用户' if msg.get('role') == 'user' else 'AI助手'
                content = msg.get('content', '')

                role_para = doc.add_paragraph()
                role_run = role_para.add_run(f'{role}：')
                role_run.bold = True

                doc.add_paragraph(content[:500])  # 限制长度
                doc.add_paragraph()

        # 页脚
        doc.add_page_break()
        footer_para = doc.add_paragraph('本报告由 AI 智能简历分析系统生成')
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 保存到内存
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer


class ResumeDocumentGenerator:
    """优化后的简历文档生成器"""

    def __init__(self):
        pass

    def _setup_styles(self, doc: Document):
        """设置文档样式"""
        # 正文样式
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Microsoft YaHei'  # 使用微软雅黑
        normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        normal_style.font.size = Pt(10.5)
        normal_style.paragraph_format.line_spacing = 1.15
        normal_style.paragraph_format.space_after = Pt(6)

        # 一级标题样式
        heading1_style = doc.styles['Heading 1']
        heading1_style.font.name = 'Microsoft YaHei'
        heading1_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        heading1_style.font.size = Pt(14)
        heading1_style.font.bold = True
        heading1_style.font.color.rgb = RGBColor(31, 78, 121)  # 深蓝色
        heading1_style.paragraph_format.space_before = Pt(12)
        heading1_style.paragraph_format.space_after = Pt(6)
        heading1_style.paragraph_format.left_indent = Pt(0)

    def _add_section_heading(self, doc: Document, text: str, icon: str = ""):
        """添加带样式的章节标题"""
        heading = doc.add_heading(level=1)

        # 不再添加图标，避免兼容性问题

        # 添加标题文字
        title_run = heading.add_run(text)
        title_run.font.name = 'Microsoft YaHei'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(31, 78, 121)

        # 添加下划线效果
        heading.paragraph_format.space_after = Pt(8)
        heading.paragraph_format.space_before = Pt(12)

    def _add_info_table(self, doc: Document, data: Dict[str, str], cols: int = 2):
        """添加信息表格"""
        rows_needed = (len(data) + cols - 1) // cols
        table = doc.add_table(rows=rows_needed, cols=cols * 2)
        table.style = 'Light Grid Accent 1'

        items = list(data.items())
        for i, (key, value) in enumerate(items):
            row_idx = i // cols
            col_idx = (i % cols) * 2

            # 设置标签单元格
            label_cell = table.cell(row_idx, col_idx)
            label_cell.text = key
            label_para = label_cell.paragraphs[0]
            label_run = label_para.runs[0]
            label_run.font.bold = True
            label_run.font.size = Pt(10)
            label_cell.width = Inches(1.2)

            # 设置值单元格
            value_cell = table.cell(row_idx, col_idx + 1)
            value_cell.text = value
            value_para = value_cell.paragraphs[0]
            value_para.runs[0].font.size = Pt(10)

        return table

    def generate_resume_word(
        self,
        resume_data: Dict[str, Any],
        template_style: str = "modern"
    ) -> BytesIO:
        """
        生成 Word 格式的简历

        Args:
            resume_data: 简历数据（结构化）
            template_style: 模板风格（modern/classic/simple）

        Returns:
            BytesIO: Word 文档的字节流
        """
        doc = Document()

        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.9)
            section.right_margin = Inches(0.9)

        # 设置样式
        self._setup_styles(doc)

        # ==================== 1. 个人信息头部 ====================
        name = resume_data.get('personal_info', {}).get('name', '求职者姓名')

        # 姓名 - 大号加粗居中
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(name)
        name_run.font.name = 'Microsoft YaHei'
        name_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        name_run.font.size = Pt(22)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(31, 78, 121)
        name_para.paragraph_format.space_after = Pt(4)

        # 联系方式 - 小号居中，分隔符美化
        contact_info = resume_data.get('personal_info', {})
        contact_items = []
        if contact_info.get('phone'):
            contact_items.append(f"电话：{contact_info['phone']}")
        if contact_info.get('email'):
            contact_items.append(f"邮箱：{contact_info['email']}")
        if contact_info.get('location'):
            contact_items.append(f"地址：{contact_info['location']}")

        if contact_items:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run('  |  '.join(contact_items))
            contact_run.font.name = 'Microsoft YaHei'
            contact_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            contact_run.font.size = Pt(10)
            contact_run.font.color.rgb = RGBColor(89, 89, 89)
            contact_para.paragraph_format.space_after = Pt(16)

        # 添加分隔线
        doc.add_paragraph('_' * 80).paragraph_format.space_after = Pt(12)

        # ==================== 2. 求职意向 ====================
        job_intent = resume_data.get('job_intent', {})
        if job_intent and any(job_intent.values()):
            self._add_section_heading(doc, '求职意向')

            intent_data = {}
            if job_intent.get('position'):
                intent_data['期望职位'] = job_intent['position']
            if job_intent.get('salary'):
                intent_data['期望薪资'] = job_intent['salary']
            if job_intent.get('location'):
                intent_data['工作地点'] = job_intent['location']

            if intent_data:
                self._add_info_table(doc, intent_data, cols=3)
                doc.add_paragraph()

        # ==================== 3. 教育背景 ====================
        education = resume_data.get('education', [])
        if education:
            self._add_section_heading(doc, '教育背景')

            for edu in education:
                # 学校 | 专业 | 学历
                edu_para = doc.add_paragraph()
                edu_para.paragraph_format.left_indent = Inches(0.2)
                edu_para.paragraph_format.space_after = Pt(2)

                school = edu.get('school', '')
                major = edu.get('major', '')
                degree = edu.get('degree', '')

                school_run = edu_para.add_run(f"• {school}")
                school_run.font.bold = True
                school_run.font.size = Pt(11)
                school_run.font.color.rgb = RGBColor(31, 78, 121)

                if major:
                    major_run = edu_para.add_run(f"  |  {major}")
                    major_run.font.size = Pt(10.5)

                if degree:
                    degree_run = edu_para.add_run(f"  |  {degree}")
                    degree_run.font.size = Pt(10.5)

                # 时间
                if edu.get('time'):
                    time_para = doc.add_paragraph()
                    time_para.paragraph_format.left_indent = Inches(0.35)
                    time_para.paragraph_format.space_after = Pt(2)
                    time_run = time_para.add_run(f"时间：{edu['time']}")
                    time_run.font.size = Pt(9.5)
                    time_run.font.color.rgb = RGBColor(128, 128, 128)

                # 描述
                if edu.get('description'):
                    desc_para = doc.add_paragraph()
                    desc_para.paragraph_format.left_indent = Inches(0.35)
                    desc_para.paragraph_format.space_after = Pt(8)
                    desc_run = desc_para.add_run(edu['description'])
                    desc_run.font.size = Pt(10)

            doc.add_paragraph()

        # ==================== 4. 专业技能 ====================
        skills = resume_data.get('skills', [])
        if skills:
            self._add_section_heading(doc, '专业技能')

            if isinstance(skills, list) and len(skills) > 0:
                if isinstance(skills[0], str):
                    # 简单列表 - 用分隔符美化
                    skills_para = doc.add_paragraph()
                    skills_para.paragraph_format.left_indent = Inches(0.2)
                    skills_run = skills_para.add_run('• ' + '  •  '.join(skills))
                    skills_run.font.size = Pt(10.5)
                else:
                    # 分类技能
                    for skill_group in skills:
                        category = skill_group.get('category', '')
                        items = skill_group.get('items', [])

                        if category and items:
                            skill_para = doc.add_paragraph()
                            skill_para.paragraph_format.left_indent = Inches(0.2)
                            skill_para.paragraph_format.space_after = Pt(4)

                            cat_run = skill_para.add_run(f"• {category}：")
                            cat_run.font.bold = True
                            cat_run.font.size = Pt(10.5)
                            cat_run.font.color.rgb = RGBColor(31, 78, 121)

                            items_run = skill_para.add_run('  •  '.join(items))
                            items_run.font.size = Pt(10.5)

            doc.add_paragraph()

        # ==================== 5. 工作经历 ====================
        work_experience = resume_data.get('work_experience', [])
        if work_experience:
            self._add_section_heading(doc, '工作经历')

            for work in work_experience:
                # 公司名称和职位
                company_para = doc.add_paragraph()
                company_para.paragraph_format.left_indent = Inches(0.2)
                company_para.paragraph_format.space_after = Pt(2)

                company_run = company_para.add_run(f"• {work.get('company', '')}")
                company_run.font.bold = True
                company_run.font.size = Pt(11)
                company_run.font.color.rgb = RGBColor(31, 78, 121)

                position = work.get('position', '')
                if position:
                    pos_run = company_para.add_run(f"  |  {position}")
                    pos_run.font.size = Pt(10.5)
                    pos_run.font.bold = True

                # 时间
                if work.get('time'):
                    time_para = doc.add_paragraph()
                    time_para.paragraph_format.left_indent = Inches(0.35)
                    time_para.paragraph_format.space_after = Pt(4)
                    time_run = time_para.add_run(f"时间：{work['time']}")
                    time_run.font.size = Pt(9.5)
                    time_run.font.color.rgb = RGBColor(128, 128, 128)

                # 工作职责
                if work.get('responsibilities'):
                    for resp in work['responsibilities']:
                        resp_para = doc.add_paragraph()
                        resp_para.paragraph_format.left_indent = Inches(0.35)
                        resp_para.paragraph_format.first_line_indent = Inches(0)
                        resp_para.paragraph_format.space_after = Pt(3)
                        resp_run = resp_para.add_run(f"  ▪ {resp}")
                        resp_run.font.size = Pt(10)

                doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # ==================== 6. 项目经验 ====================
        projects = resume_data.get('projects', [])
        if projects:
            self._add_section_heading(doc, '项目经验')

            for project in projects:
                # 项目名称
                proj_para = doc.add_paragraph()
                proj_para.paragraph_format.left_indent = Inches(0.2)
                proj_para.paragraph_format.space_after = Pt(2)

                proj_run = proj_para.add_run(f"• {project.get('name', '')}")
                proj_run.font.bold = True
                proj_run.font.size = Pt(11)
                proj_run.font.color.rgb = RGBColor(31, 78, 121)

                # 时间
                if project.get('time'):
                    time_run = proj_para.add_run(f"  ({project['time']})")
                    time_run.font.size = Pt(9.5)
                    time_run.font.color.rgb = RGBColor(128, 128, 128)

                # 项目描述
                if project.get('description'):
                    desc_para = doc.add_paragraph()
                    desc_para.paragraph_format.left_indent = Inches(0.35)
                    desc_para.paragraph_format.space_after = Pt(3)
                    desc_run = desc_para.add_run(project['description'])
                    desc_run.font.size = Pt(10)

                # 技术栈
                if project.get('tech_stack'):
                    tech_para = doc.add_paragraph()
                    tech_para.paragraph_format.left_indent = Inches(0.35)
                    tech_para.paragraph_format.space_after = Pt(3)

                    tech_label = tech_para.add_run('技术栈：')
                    tech_label.font.bold = True
                    tech_label.font.size = Pt(10)

                    tech_items = tech_para.add_run('  •  '.join(project['tech_stack']))
                    tech_items.font.size = Pt(10)

                # 项目成果
                if project.get('achievements'):
                    for achievement in project['achievements']:
                        ach_para = doc.add_paragraph()
                        ach_para.paragraph_format.left_indent = Inches(0.35)
                        ach_para.paragraph_format.space_after = Pt(3)
                        ach_run = ach_para.add_run(f"  • {achievement}")
                        ach_run.font.size = Pt(10)
                        # 使用默认黑色，不设置绿色

                doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # ==================== 7. 个人评价 ====================
        self_evaluation = resume_data.get('self_evaluation', '')
        if self_evaluation:
            self._add_section_heading(doc, '个人评价')

            eval_para = doc.add_paragraph()
            eval_para.paragraph_format.left_indent = Inches(0.2)
            eval_para.paragraph_format.first_line_indent = Inches(0.3)
            eval_run = eval_para.add_run(self_evaluation)
            eval_run.font.size = Pt(10.5)

        # 保存到内存
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer

    def generate_simple_resume_word(
        self,
        analysis_data: Dict[str, Any],
        optimized_content: str = None
    ) -> BytesIO:
        """
        基于分析数据生成简单格式的简历（使用专业样式）

        Args:
            analysis_data: 简历分析数据
            optimized_content: AI 优化后的完整文本（可选）

        Returns:
            BytesIO: Word 文档的字节流
        """
        doc = Document()

        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.9)
            section.right_margin = Inches(0.9)

        # 设置样式
        self._setup_styles(doc)

        # ==================== 1. 标题 ====================
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run('个人简历')
        name_run.font.name = 'Microsoft YaHei'
        name_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        name_run.font.size = Pt(22)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(31, 78, 121)
        name_para.paragraph_format.space_after = Pt(16)

        # 添加分隔线
        doc.add_paragraph('_' * 80).paragraph_format.space_after = Pt(12)

        # ==================== 2. 基本信息 ====================
        self._add_section_heading(doc, '基本信息')

        info_data = {
            '学历': analysis_data.get('education', '未提供'),
            '工作年限': f"{analysis_data.get('experience_years', 0)} 年"
        }
        self._add_info_table(doc, info_data, cols=2)
        doc.add_paragraph()

        # ==================== 3. 求职意向 ====================
        if analysis_data.get('desired_position') or analysis_data.get('desired_salary'):
            self._add_section_heading(doc, '求职意向')

            intent_data = {}
            if analysis_data.get('desired_position'):
                intent_data['期望职位'] = analysis_data['desired_position']
            if analysis_data.get('desired_salary'):
                intent_data['期望薪资'] = analysis_data['desired_salary']

            if intent_data:
                self._add_info_table(doc, intent_data, cols=2)
                doc.add_paragraph()

        # ==================== 4. 专业技能 ====================
        skills = analysis_data.get('skills', [])
        if skills:
            self._add_section_heading(doc, '专业技能')

            skills_para = doc.add_paragraph()
            skills_para.paragraph_format.left_indent = Inches(0.2)
            skills_run = skills_para.add_run('• ' + '  •  '.join(skills))
            skills_run.font.size = Pt(10.5)
            doc.add_paragraph()

        # ==================== 5. 核心优势 ====================
        key_strengths = analysis_data.get('key_strengths', [])
        if key_strengths:
            self._add_section_heading(doc, '核心优势')

            for i, strength in enumerate(key_strengths, 1):
                strength_para = doc.add_paragraph()
                strength_para.paragraph_format.left_indent = Inches(0.2)
                strength_para.paragraph_format.space_after = Pt(4)

                num_run = strength_para.add_run(f"{i}. ")
                num_run.font.bold = True
                num_run.font.size = Pt(10.5)
                num_run.font.color.rgb = RGBColor(31, 78, 121)

                text_run = strength_para.add_run(strength)
                text_run.font.size = Pt(10.5)

            doc.add_paragraph()

        # ==================== 6. 工作经历 ====================
        work_exp = analysis_data.get('work_experience', '')
        if work_exp:
            self._add_section_heading(doc, '工作经历')

            # 分段显示工作经历
            paragraphs = work_exp.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    work_para = doc.add_paragraph()
                    work_para.paragraph_format.left_indent = Inches(0.2)
                    work_para.paragraph_format.space_after = Pt(6)
                    work_run = work_para.add_run(para_text.strip())
                    work_run.font.size = Pt(10)

            doc.add_paragraph()

        # ==================== 7. AI 优化建议（如果有） ====================
        if optimized_content:
            doc.add_page_break()
            self._add_section_heading(doc, 'AI 优化建议')

            opt_para = doc.add_paragraph()
            opt_para.paragraph_format.left_indent = Inches(0.2)
            opt_para.paragraph_format.first_line_indent = Inches(0.3)
            opt_run = opt_para.add_run(optimized_content)
            opt_run.font.size = Pt(10.5)

        # 保存到内存
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer
